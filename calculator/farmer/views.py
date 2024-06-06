import json
import random
from datetime import datetime, timedelta
from itertools import groupby
from operator import attrgetter

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.utils.dateformat import DateFormat
from django.utils.timezone import make_aware, is_naive
from django.views import View
from docx import Document
from django.db.models import Count
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.views.generic import ListView, TemplateView

from .forms import ProductionPlanForm, UserCreationForm
from .models import ProductionDepartment, Product, RawMaterialStock, Stock, Recipe, ProductionPlan, ManagementOrder
from django.shortcuts import render, get_object_or_404,redirect
from django.core.exceptions import ValidationError
from django.db.models import Sum, F
from django.db import transaction
from collections import defaultdict
import os
import tempfile
from zipfile import ZipFile
import logging

logger = logging.getLogger(__name__)

@login_required
def create_user(request):
    if not request.user.is_superuser:
        return redirect('home')  # Перенаправление на главную страницу, если не суперпользователь

    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data['password'])
            user.save()
            return redirect('home')  # Перенаправление на главную страницу после успешного создания
    else:
        form = UserCreationForm()

    return render(request, 'farmer/create_user.html', {'form': form})

@login_required
@require_POST
def import_management_orders(request):
    try:
        json_file = request.FILES.get('json_file')
        if not json_file.name.endswith('.json'):
            return JsonResponse({'error': 'Неверный формат файла. Требуется JSON.'}, status=400)

        data = json.load(json_file)
        for item in data:
            try:
                # Проверка и получение продукта по коду
                product = Product.objects.get(id=item['product_code'])

                # Проверка на дубликаты перед созданием
                if ManagementOrder.objects.filter(product_code=product, start_date=item['start_date']).exists():
                    return JsonResponse({
                                            'error': f'Заявка с кодом продукта {item["product_code"]} на дату {item["start_date"]} уже существует.'},
                                        status=400)

                # Создание новой заявки
                ManagementOrder.objects.create(
                    Order_id=item['Order_id'],
                    product_code=product,
                    start_date=item['start_date'],
                    end_date=item['end_date'],
                    quantity=item['quantity']
                )
            except Product.DoesNotExist:
                return JsonResponse({'error': f'Продукт с кодом {item["product_code"]} не найден.'}, status=400)
            except (KeyError, TypeError, ValidationError) as e:
                return JsonResponse({'error': str(e)}, status=400)

        return JsonResponse({'success': 'Данные успешно импортированы.'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)



@login_required
def orders_list(request):
    orders = ManagementOrder.objects.values('Order_id').annotate(total=Count('id')).order_by('Order_id')
    return render(request, 'farmer/orders_list.html', {'orders': orders})
@login_required
def order_details(request, order_id):
    # Получаем все заказы с данным Order_id
    orders = ManagementOrder.objects.filter(Order_id=order_id).select_related('product_code')
    # Получаем все рецепты для продуктов в заказах
    recipes = Recipe.objects.filter(product__in=[order.product_code for order in orders]).select_related('raw_material')
    # Передаем order_id, коллекцию заказов и рецепты в шаблон
    return render(request, 'farmer/order_details.html', {
        'orders': orders,
        'order_id': order_id,
        'recipes': recipes,
    })
@login_required
def home(request):
    return render(request, 'farmer/home.html')
@login_required
def show_import_form(request):
    return render(request, 'farmer/import.html')


class MultiProductProductionPlanView(LoginRequiredMixin,TemplateView):
    template_name = 'farmer/multi_product_production_plan.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Получаем все продукты
        products = Product.objects.all()
        context['products'] = products

        if self.request.method == 'POST':
            # Получаем список продуктов и их параметры из POST-запроса
            products_ids = self.request.POST.getlist('products')
            start_dates = self.request.POST.getlist('start_dates')
            end_dates = self.request.POST.getlist('end_dates')
            quantities = self.request.POST.getlist('quantities')

            # Словарь для хранения общего необходимого количества каждого сырья
            total_required_materials = defaultdict(int)

            orders = []
            for i, product_id in enumerate(products_ids):
                product = Product.objects.get(id=product_id)
                required_quantity = int(quantities[i])
                start_date = datetime.strptime(start_dates[i], '%Y-%m-%d')
                end_date = datetime.strptime(end_dates[i], '%Y-%m-%d')

                order = {
                    'product': product,
                    'start_date': start_date,
                    'end_date': end_date,
                    'quantity': required_quantity
                }
                orders.append(order)

                recipes = Recipe.objects.filter(product=product)
                for recipe in recipes:
                    total_required_materials[recipe.raw_material.name] += recipe.required_quantity * required_quantity

            # Используем транзакцию для проверки наличия сырья и возможных операций с базой данных
            with transaction.atomic():
                # Словарь для хранения информации о возможности производства для каждого заказа
                production_possibility = {}
                raw_materials_stock = RawMaterialStock.objects.filter(name__in=total_required_materials.keys())
                stock_dict = {stock.name: stock.quantity for stock in raw_materials_stock}

                # Проверяем наличие сырья на складе и корректируем остатки по мере использования
                for order in orders:
                    product = order['product']
                    required_quantity = order['quantity']
                    recipes = Recipe.objects.filter(product=product)
                    sufficient_resources = True
                    shortages = {}

                    for recipe in recipes:
                        material_name = recipe.raw_material.name
                        required_amount = recipe.required_quantity * required_quantity
                        available_quantity = stock_dict.get(material_name, 0)

                        if available_quantity < required_amount:
                            shortages[material_name] = required_amount - available_quantity
                            sufficient_resources = False
                        else:
                            # Временно вычитаем необходимое количество из доступного
                            stock_dict[material_name] -= required_amount

                    # Проверяем возможность выполнения заказов
                    if sufficient_resources:
                        production_department = product.production_department
                        monthly_output = production_department.average_output

                        # Определяем текущую загруженность производства
                        overlapping_orders = ProductionPlan.objects.filter(
                            product__production_department=production_department,
                            order__end_date__gte=order['start_date'],
                            order__start_date__lte=order['end_date']
                        )

                        current_load = 0
                        for overlapping_order in overlapping_orders:
                            overlapping_order_start = overlapping_order.order.start_date
                            if isinstance(overlapping_order_start, datetime):
                                overlapping_order_start = overlapping_order_start.date()
                            overlapping_order_end = overlapping_order.order.end_date
                            if isinstance(overlapping_order_end, datetime):
                                overlapping_order_end = overlapping_order_end.date()
                            order_start = order['start_date'].date() if isinstance(order['start_date'], datetime) else order['start_date']
                            order_end = order['end_date'].date() if isinstance(order['end_date'], datetime) else order['end_date']

                            overlap_days = (min(overlapping_order_end, order_end) - max(overlapping_order_start, order_start)).days
                            if overlap_days > 0:
                                overlap_months = overlap_days / 30
                                current_load += overlapping_order.planned_quantity / overlap_months

                        available_capacity = monthly_output - current_load
                        required_months = required_quantity / monthly_output

                        if available_capacity <= 0 or required_quantity > available_capacity * required_months:
                            delay = (required_quantity / available_capacity) - required_months if available_capacity > 0 else required_months
                            if delay > 0.01:  # Установим порог для задержки
                                production_possibility[product.name] = f"Производственный отдел не сможет уложиться в сроки, просрочка составит {delay:.2f} дней"
                            else:
                                production_possibility[product.name] = f"Можно произвести {required_quantity} кг за {required_months:.2f} месяцев"
                        else:
                            production_possibility[product.name] = f"Можно произвести {required_quantity} кг за {required_months:.2f} месяцев"
                    else:
                        shortage_info = ', '.join([f"{material}: недостает {amount} ед." for material, amount in shortages.items()])
                        production_possibility[product.name] = f"Недостаточно сырья: {shortage_info}"

            # Отображаем результаты расчёта
            context['orders'] = orders
            context['production_possibility'] = production_possibility

        return context

    def post(self, request, *args, **kwargs):
        context = self.get_context_data(**kwargs)
        return self.render_to_response(context)


def calculate_multi_product_production_plan(request):
    if request.method == 'POST':
        products = request.POST.getlist('products')
        start_dates = request.POST.getlist('start_dates')
        end_dates = request.POST.getlist('end_dates')
        quantities = request.POST.getlist('quantities')

        orders = []
        for product, start_date, end_date, quantity in zip(products, start_dates, end_dates, quantities):
            orders.append({
                'product_id': int(product),
                'start_date': datetime.strptime(start_date, '%Y-%m-%d'),
                'end_date': datetime.strptime(end_date, '%Y-%m-%d'),
                'quantity': int(quantity)
            })

        production_possibility = calculate_production_possibility(orders)
        context = {
            'products': products,
            'start_dates': start_dates,
            'end_dates': end_dates,
            'quantities': quantities,
            'production_possibility': production_possibility,
            'all_products': Product.objects.all(),  # Добавьте эту строку
            'orders': orders  # Добавьте эту строку для передачи заказов в шаблон
        }
        return render(request, 'farmer/production_plan_form.html', context)


def calculate_production_possibility(orders):
    production_possibility = {}
    total_required_materials = defaultdict(int)

    for order in orders:
        product = Product.objects.get(id=order['product_id'])
        required_quantity = order['quantity']
        recipes = Recipe.objects.filter(product=product)

        for recipe in recipes:
            total_required_materials[recipe.raw_material.name] += recipe.required_quantity * required_quantity

    raw_materials_stock = RawMaterialStock.objects.filter(name__in=total_required_materials.keys())
    stock_dict = {stock.name: stock for stock in raw_materials_stock}

    for order in orders:
        product = Product.objects.get(id=order['product_id'])
        required_quantity = order['quantity']
        recipes = Recipe.objects.filter(product=product)
        sufficient_resources = True
        shortages = {}

        for recipe in recipes:
            material_name = recipe.raw_material.name
            required_amount = recipe.required_quantity * required_quantity
            available_quantity = stock_dict[material_name].quantity

            if available_quantity < required_amount:
                shortages[material_name] = required_amount - available_quantity
                sufficient_resources = False
            else:
                stock_dict[material_name].quantity -= required_amount

        if sufficient_resources:
            production_department = product.production_department
            production_time = (order['end_date'] - order['start_date']).days
            monthly_output = production_department.average_output
            required_months = required_quantity / monthly_output

            overlapping_orders = ProductionPlan.objects.filter(
                product__production_department=production_department,
                order__end_date__gte=datetime.combine(order['start_date'], datetime.min.time()),  # Преобразование в datetime.datetime
                order__start_date__lte=datetime.combine(order['end_date'], datetime.min.time())  # Преобразование в datetime.datetime
            )

            current_load = 0
            for overlapping_order in overlapping_orders:
                overlap_days = (min(overlapping_order.order.end_date, order['end_date']) - max(
                    overlapping_order.order.start_date, order['start_date'])).days
                overlap_months = overlap_days / 30
                current_load += overlapping_order.planned_quantity / overlap_months

            available_capacity = monthly_output - current_load
            if available_capacity <= 0 or required_quantity > available_capacity * required_months:
                delay = (
                    required_quantity / available_capacity) - required_months if available_capacity > 0 else required_months
                if delay > 0:
                    production_possibility[
                        product.name] = f"Производственный отдел не сможет уложиться в сроки, просрочка составит {delay:.2f} дней"
                else:
                    production_possibility[
                        product.name] = f"Можно произвести {required_quantity} единиц за {required_months:.2f} месяцев"
        else:
            shortage_info = ', '.join(
                [f"{material}: недостает {amount} ед." for material, amount in shortages.items()])
            production_possibility[product.name] = f"Недостаточно сырья: {shortage_info}"

    return production_possibility

@login_required
def calculate_production_plan(request, order_id):
    # Получаем все заявки с данным order_id
    orders = ManagementOrder.objects.filter(Order_id=order_id).select_related('product_code',
                                                                              'product_code__production_department')

    # Словарь для хранения общего необходимого количества каждого сырья
    total_required_materials = defaultdict(int)

    # Собираем информацию о необходимом количестве сырья для всех продуктов
    for order in orders:
        product = order.product_code
        required_quantity = order.quantity
        recipes = Recipe.objects.filter(product=product)

        for recipe in recipes:
            total_required_materials[recipe.raw_material.name] += recipe.required_quantity * required_quantity

    # Используем транзакцию для проверки наличия сырья и возможных операций с базой данных
    with transaction.atomic():
        # Словарь для хранения информации о возможности производства для каждого заказа
        production_possibility = {}
        materials_shortage = {}
        raw_materials_stock = RawMaterialStock.objects.filter(name__in=total_required_materials.keys())
        stock_dict = {stock.name: stock.quantity for stock in raw_materials_stock}

        # Проверяем наличие сырья на складе и корректируем остатки по мере использования
        for order in orders:
            product = order.product_code
            required_quantity = order.quantity
            recipes = Recipe.objects.filter(product=product)
            sufficient_resources = True
            shortages = {}

            for recipe in recipes:
                material_name = recipe.raw_material.name
                required_amount = recipe.required_quantity * required_quantity
                available_quantity = stock_dict.get(material_name, 0)

                if available_quantity < required_amount:
                    shortages[material_name] = required_amount - available_quantity
                    sufficient_resources = False
                else:
                    # Временно вычитаем необходимое количество из доступного
                    stock_dict[material_name] -= required_amount

            # Проверяем возможность выполнения заказов
            if sufficient_resources:
                production_department = product.production_department
                production_time = (order.end_date - order.start_date).days
                monthly_output = production_department.average_output
                required_months = required_quantity / monthly_output

                if production_time < required_months * 30:
                    delay = (required_months * 30) - production_time
                    production_possibility[
                        product.name] = f"Производственный отдел не сможет уложиться в сроки, просрочка составит {delay:.2f} дней"
                else:
                    production_possibility[
                        product.name] = f"Можно произвести {required_quantity} единиц за {required_months:.2f} месяцев"
            else:
                shortage_info = ', '.join(
                    [f"{material}: недостает {amount} ед." for material, amount in shortages.items()])
                production_possibility[product.name] = f"Недостаточно сырья: {shortage_info}"

    # Отображаем результаты расчёта
    context = {
        'orders': orders,
        'production_possibility': production_possibility,
    }
    return render(request, 'farmer/production_plan.html', context)

@login_required
def success_page(request):
    return render(request, 'farmer/success_page.html')

def save_multi_product_production_plan(request):
    if request.method == 'POST':
        products = request.POST.getlist('products')
        start_dates = request.POST.getlist('start_dates')
        end_dates = request.POST.getlist('end_dates')
        quantities = request.POST.getlist('quantities')

        orders = []
        for product, start_date, end_date, quantity in zip(products, start_dates, end_dates, quantities):
            orders.append({
                'product_id': int(product),
                'start_date': datetime.strptime(start_date, '%Y-%m-%d'),
                'end_date': datetime.strptime(end_date, '%Y-%m-%d'),
                'quantity': int(quantity)
            })

        production_plans = []
        manager_report = []

        total_required_materials = defaultdict(int)
        for order in orders:
            product = Product.objects.get(id=order['product_id'])
            required_quantity = order['quantity']
            recipes = Recipe.objects.filter(product=product)
            for recipe in recipes:
                total_required_materials[recipe.raw_material.name] += recipe.required_quantity * required_quantity

        raw_materials_stock = RawMaterialStock.objects.filter(name__in=total_required_materials.keys())
        stock_dict = {stock.name: stock for stock in raw_materials_stock}

        for order in orders:
            product = Product.objects.get(id=order['product_id'])
            required_quantity = order['quantity']
            recipes = Recipe.objects.filter(product=product)
            sufficient_resources = True
            shortages = {}

            for recipe in recipes:
                material_name = recipe.raw_material.name
                required_amount = recipe.required_quantity * required_quantity
                available_quantity = stock_dict[material_name].quantity

                if available_quantity < required_amount:
                    shortages[material_name] = required_amount - available_quantity
                    sufficient_resources = False
                else:
                    stock_dict[material_name].quantity -= required_amount

            if sufficient_resources:
                production_department = product.production_department
                production_time = (order['end_date'] - order['start_date']).days
                monthly_output = production_department.average_output
                required_months = required_quantity / monthly_output

                overlapping_orders = ProductionPlan.objects.filter(
                    product__production_department=production_department,
                    order__end_date__gte=order['start_date'],
                    order__start_date__lte=order['end_date']
                )

                current_load = 0
                for overlapping_order in overlapping_orders:
                    overlap_days = (min(overlapping_order.order.end_date, order['end_date']) - max(
                        overlapping_order.order.start_date, order['start_date'])).days
                    overlap_months = overlap_days / 30
                    current_load += overlapping_order.planned_quantity / overlap_months

                available_capacity = monthly_output - current_load
                if available_capacity <= 0 or required_quantity > available_capacity * required_months:
                    delay = (
                                    required_quantity / available_capacity) - required_months if available_capacity > 0 else required_months
                    if delay > 0:
                        manager_report.append(
                            f"Производственный отдел {product.name} не сможет уложиться в сроки, просрочка составит {delay:.2f} дней")
                    else:
                        manager_report.append(
                            f"Производственный отдел {product.name} может произвести {required_quantity} единиц за {required_months:.2f} месяцев")
                else:
                    manager_report.append(
                        f"Производственный отдел {product.name} может произвести {required_quantity} единиц за {required_months:.2f} месяцев")

                    # Создаем заказ и план производства
                    management_order = ManagementOrder.objects.create(
                        product=product,
                        start_date=order['start_date'],
                        end_date=order['end_date'],
                        quantity=required_quantity
                    )

                    production_plan = ProductionPlan.objects.create(
                        order=management_order,
                        planned_quantity=required_quantity,
                        actual_quantity=0,  # изначально фактическое количество 0
                    )

                    production_plans.append(production_plan)
            else:
                shortage_info = ', '.join(
                    [f"{material}: недостает {amount} ед." for material, amount in shortages.items()])
                manager_report.append(f"Недостаточно сырья для производства {product.name}: {shortage_info}")

        # Создаем временный файл ZIP для отчетов
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_filename = os.path.join(temp_dir, 'production_plans.zip')
            with ZipFile(zip_filename, 'w') as zip_file:
                for plan in production_plans:
                    filename = f"ProductionPlan_{plan.order.product.name}_{plan.order.start_date}.txt"
                    filepath = os.path.join(temp_dir, filename)
                    with open(filepath, 'w') as file:
                        file.write(f"Product: {plan.order.product.name}\n")
                        file.write(f"Start Date: {plan.order.start_date}\n")
                        file.write(f"End Date: {plan.order.end_date}\n")
                        file.write(f"Planned Quantity: {plan.planned_quantity}\n")
                        file.write(f"Actual Quantity: {plan.actual_quantity}\n")
                    zip_file.write(filepath, filename)

            # Отправляем ZIP файл в ответ
            with open(zip_filename, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/zip')
                response['Content-Disposition'] = f'attachment; filename="production_plans.zip"'
                return response
    else:
        return HttpResponse("Invalid request method", status=405)


class SaveProductionPlanView(LoginRequiredMixin,View):
    def post(self, request):
        # Получаем данные из POST-запроса
        print(request.POST)
        products = request.POST.getlist('product')
        start_dates = request.POST.getlist('start_date')
        end_dates = request.POST.getlist('end_date')
        quantities = request.POST.getlist('quantity')

        if products and start_dates and end_dates and quantities:
                # Создаем новый заказ
            order_id = generate_order_id()


            # Обновляем количество продуктов на складе и создаем производственные планы
            for product_id, start_date, end_date, quantity in zip(products, start_dates, end_dates, quantities):


                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()


                new_order = ManagementOrder.objects.create(Order_id=order_id,
                                                           product_code_id=product_id,
                                                           start_date=start_date,
                                                           end_date=end_date,
                                                           quantity=quantity)  # Указываем product_code
                product = Product.objects.get(id=product_id)
                # Получаем рецепты для данного продукта
                recipes = Recipe.objects.filter(product=product)
                # Обновляем количество сырья на складе
                for recipe in recipes:
                    raw_material = recipe.raw_material
                    required_quantity = recipe.required_quantity * quantity
                    # Проверяем наличие достаточного количества сырья на складе
                    if raw_material.quantity < required_quantity:
                        # Если не хватает сырья, выводим сообщение об ошибке
                        raise Exception(f"Недостаточно сырья {raw_material.name} на складе")
                    # Обновляем количество сырья на складе
                    raw_material.quantity -= required_quantity
                    raw_material.save()

            for product, start_date, end_date, quantity in zip(products, start_dates, end_dates, quantities):
                ProductionPlan.objects.create(
                    order=new_order,
                    product_id=product,
                    planned_quantity=quantity
                )

                    # Перенаправляем пользователя после сохранения
            return redirect('farmer:production_plan_list')
        else:
            # Вывести сообщение об ошибке или выполнить другое действие в случае пустых списков
            return HttpResponse("Ошибка: Не удалось получить данные из формы")

def generate_order_id():
    prefix = 'PROIS'
    # Генерируем уникальный идентификатор, например, добавляя случайное число к префиксу
    unique_id = random.randint(10000, 99999)
    return unique_id
@login_required
def save_production_plan(request, order_id):
    if request.method == 'POST':
        # Получаем все заявки с данным order_id
        orders = ManagementOrder.objects.filter(Order_id=order_id).select_related('product_code',
                                                                                  'product_code__production_department')

        # Словарь для хранения информации о возможности производства для каждого заказа
        production_possibility = {}
        production_plans = []
        manager_report = []

        # Словарь для хранения общего необходимого количества каждого сырья
        total_required_materials = defaultdict(int)
        for order in orders:
            product = order.product_code
            required_quantity = order.quantity
            recipes = Recipe.objects.filter(product=product)
            for recipe in recipes:
                total_required_materials[recipe.raw_material.name] += recipe.required_quantity * required_quantity

        raw_materials_stock = RawMaterialStock.objects.filter(name__in=total_required_materials.keys())
        stock_dict = {stock.name: stock for stock in raw_materials_stock}

        for order in orders:
            product = order.product_code
            required_quantity = order.quantity
            recipes = Recipe.objects.filter(product=product)
            sufficient_resources = True
            shortages = {}

            for recipe in recipes:
                material_name = recipe.raw_material.name
                required_amount = recipe.required_quantity * required_quantity
                available_quantity = stock_dict[material_name].quantity

                if available_quantity < required_amount:
                    shortages[material_name] = required_amount - available_quantity
                    sufficient_resources = False
                else:
                    stock_dict[material_name].quantity -= required_amount

            if sufficient_resources:
                production_department = product.production_department
                production_time = (order.end_date - order.start_date).days
                monthly_output = production_department.average_output
                required_months = required_quantity / monthly_output

                if production_time < required_months * 30:
                    delay = (required_months * 30) - production_time
                    production_possibility[
                        product.name] = f"Производственный отдел не сможет уложиться в сроки, просрочка составит {delay:.2f} дней"
                    manager_report.append((product.name,
                                           f"Производственный отдел не сможет уложиться в сроки, просрочка составит {delay:.2f} дней"))
                else:
                    production_possibility[
                        product.name] = f"Можно произвести {required_quantity} единиц за {required_months:.2f} месяцев"
                    production_plans.append(
                        ProductionPlan(order=order, product=product, planned_quantity=required_quantity))
            else:
                shortage_info = ', '.join(
                    [f"{material}: недостает {amount} ед." for material, amount in shortages.items()])
                production_possibility[product.name] = f"Недостаточно сырья: {shortage_info}"
                manager_report.append((product.name, f"Недостаточно сырья: {shortage_info}"))

        # Обновляем количество сырья на складе
        for stock in stock_dict.values():
            stock.save()

        # Сохраняем производственный план в базе данных
        ProductionPlan.objects.bulk_create(production_plans)

        # Генерируем документы
        temp_dir = tempfile.mkdtemp()
        department_files = generate_department_documents(production_plans, temp_dir)
        manager_report_file = generate_manager_report(manager_report, temp_dir)

        # Создаем zip-архив с документами
        zip_path = os.path.join(temp_dir, 'production_plan_documents.zip')
        with ZipFile(zip_path, 'w') as zip_file:
            for file_path in department_files:
                zip_file.write(file_path, os.path.basename(file_path))
            zip_file.write(manager_report_file, os.path.basename(manager_report_file))

        # Отправляем архив пользователю
        with open(zip_path, 'rb') as zip_file:
            response = HttpResponse(zip_file.read(), content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename=production_plan_documents.zip'
            return response

def generate_department_documents(production_plans, temp_dir):
    department_plans = defaultdict(list)
    for plan in production_plans:
        department_plans[plan.product.production_department.name].append(plan)

    file_paths = []
    for department, plans in department_plans.items():
        document = Document()
        document.add_heading(f'План производства для отдела {department}', 0)

        for plan in plans:
            document.add_heading(plan.product.name, level=1)
            document.add_paragraph(f'Запланированное количество: {plan.planned_quantity}')
            document.add_paragraph(f'Дата начала: {plan.order.start_date}')
            document.add_paragraph(f'Дата окончания: {plan.order.end_date}')

        file_path = os.path.join(temp_dir, f'{department}_production_plan.docx')
        document.save(file_path)
        file_paths.append(file_path)

    return file_paths

def generate_manager_report(manager_report, temp_dir):
    document = Document()
    document.add_heading('Отчет для менеджеров', 0)

    for product_name, issue in manager_report:
        document.add_heading(product_name, level=1)
        document.add_paragraph(issue)

    file_path = os.path.join(temp_dir, 'manager_report.docx')
    document.save(file_path)
    return file_path

class ProductionDepartmentsListView(LoginRequiredMixin,ListView):
    model = ProductionDepartment
    template_name = 'farmer/production_departments_list.html'
    context_object_name = 'departments'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Получаем выбранный месяц из GET параметров
        month_str = self.request.GET.get('month')
        if month_str:
            selected_month = datetime.strptime(month_str, '%Y-%m')
        else:
            selected_month = datetime.now()

        # Определяем временной интервал для выбранного месяца
        start_date = datetime(selected_month.year, selected_month.month, 1)
        end_date = (start_date + timedelta(days=31)).replace(day=1) - timedelta(seconds=1)

        # Преобразование в осознанное время только если время наивное
        if is_naive(start_date):
            start_date = make_aware(start_date)
        if is_naive(end_date):
            end_date = make_aware(end_date)

        # Получение только утвержденных заказов, которые включены в производственный план в выбранный месяц
        production_plans = ProductionPlan.objects.select_related('order', 'product__production_department').filter(
            order__start_date__lte=end_date,
            order__end_date__gte=start_date
        )

        # Группировка заказов по производственным отделам
        departments = ProductionDepartment.objects.all()
        load_data = []
        for department in departments:
            department_orders = []
            for plan in production_plans:
                if plan.product.production_department == department:
                    department_orders.append({
                        'order_id': plan.order.Order_id,
                        'product': plan.product.name,
                        'start_date': plan.order.start_date.strftime('%Y-%m-%d'),
                        'end_date': plan.order.end_date.strftime('%Y-%m-%d'),
                        'quantity': plan.order.quantity,
                        'monthly_output': department.average_output
                    })
            load_data.append({
                'department': department.name,
                'load': department_orders
            })

        context['load_data'] = load_data
        context['start_date'] = start_date.strftime('%Y-%m-%d')
        context['end_date'] = end_date.strftime('%Y-%m-%d')
        context['selected_month'] = selected_month.strftime('%Y-%m')
        return context

class ProductsListView(LoginRequiredMixin,ListView):
    model = Product
    template_name = 'farmer/products_list.html'
    context_object_name = 'products'

class RawMaterialStockListView(LoginRequiredMixin,ListView):
    model = RawMaterialStock
    template_name = 'farmer/raw_material_stock_list.html'
    context_object_name = 'raw_materials'

class StockListView(LoginRequiredMixin,ListView):
    model = Stock
    template_name = 'farmer/stock_list.html'
    context_object_name = 'stocks'

@login_required
def recipes_list(request):
    # Получаем все рецепты
    recipes = Recipe.objects.select_related('product', 'raw_material')

    # Группируем рецепты по продукту
    grouped_recipes = {}
    for recipe in recipes:
        if recipe.product.name not in grouped_recipes:
            grouped_recipes[recipe.product.name] = []
        grouped_recipes[recipe.product.name].append(recipe)

    # Передаем сгруппированные рецепты в шаблон
    return render(request, 'farmer/recipes_list.html', {
        'grouped_recipes': grouped_recipes,
    })

class RecipesListView(LoginRequiredMixin,ListView):
    model = Recipe
    template_name = 'farmer/recipes_list.html'
    context_object_name = 'recipes'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        recipes = self.get_queryset()

        grouped_recipes = defaultdict(list)
        for recipe in recipes:
            grouped_recipes[recipe.product.name].append(recipe)

        context['grouped_recipes'] = grouped_recipes
        return context

class ProductionPlansListView(LoginRequiredMixin,ListView):
    model = ProductionPlan
    template_name = 'farmer/production_plans_list.html'
    context_object_name = 'plans'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        plans = self.get_queryset().select_related('order', 'product').order_by('order__Order_id')
        grouped_plans = groupby(plans, key=attrgetter('order'))
        grouped_plans_list = [(order, list(plans)) for order, plans in grouped_plans]
        context['grouped_plans'] = grouped_plans_list
        return context
